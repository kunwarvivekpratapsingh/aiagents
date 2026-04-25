"""DOCX loader — extracts text from Word documents using python-docx."""
from __future__ import annotations

import logging
from pathlib import Path

from .base import Document

logger = logging.getLogger(__name__)


class DocxLoader:
    def load(self, path: Path) -> list[Document]:
        try:
            from docx import Document as DocxDocument
        except ImportError as exc:
            raise ImportError("Install python-docx: pip install python-docx") from exc

        doc = DocxDocument(str(path))

        # Collect paragraphs; preserve heading structure via double newlines
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        # Also pull text from tables
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))

        text = "\n\n".join(paragraphs)
        if not text.strip():
            logger.warning("No text extracted from DOCX: %s", path)
            return []

        return [
            Document(
                content=text,
                metadata={
                    "source": str(path.resolve()),
                    "filename": path.name,
                    "filetype": "docx",
                },
            )
        ]
